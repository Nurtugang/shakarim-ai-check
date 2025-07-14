import os
import logging

import tempfile

import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .list_utils import get_list_info

from pdf2docx import Converter

from django.core.files.base import ContentFile
from django.core.files.storage import default_storage

logger = logging.getLogger(__name__)

def process_document(file):
    """ Обрабатывает загруженный документ и извлекает из него текст и HTML """
    file_ext = get_file_extension(file.name)
    
    if file_ext == 'pdf':
        plain_text = extract_text_from_pdf(file)
        file.seek(0)  # Сбрасываем указатель для HTML
        html_content = extract_html_from_pdf(file)
        return plain_text, html_content
    elif file_ext in ['docx', 'doc']:
        plain_text = extract_text_from_docx(file)
        file.seek(0)  # Сбрасываем указатель для HTML
        html_content = extract_html_from_docx(file)
        return plain_text, html_content
    else:
        raise ValueError(f"Неподдерживаемый формат файла: {file_ext}. Поддерживаются только PDF и DOCX.")

def get_file_extension(file_name):
    _, extension = os.path.splitext(file_name)
    return extension.lower()[1:]

def save_uploaded_file(uploaded_file, user_id):
    file_name = uploaded_file.name
    file_path = f'uploads/{user_id}/{file_name}'
    path = default_storage.save(file_path, ContentFile(uploaded_file.read()))
    return path

def is_scanned_pdf(text):
    if not text or len(text.strip()) < 50:
        return True
    if text.count('\n') > len(text) * 0.5:
        return True
    special_chars = sum(1 for c in text if not c.isalnum() and not c.isspace())
    if special_chars > len(text) * 0.3:
        return True
    return False

def extract_html_from_pdf(file):
    try:
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as pdf_temp:
            for chunk in file.chunks():
                pdf_temp.write(chunk)
            pdf_path = pdf_temp.name
            
        docx_path = pdf_path.replace('.pdf', '.docx')
        
        try:
            cv = Converter(pdf_path)
            cv.convert(docx_path)
            cv.close()
            
            with open(docx_path, 'rb') as docx_file:
                return extract_html_from_docx(docx_file)
                
        finally:
            try:
                os.unlink(pdf_path)
                os.unlink(docx_path)
            except Exception as e:
                logger.warning(f"Failed to delete temporary files: {str(e)}")
                
    except Exception as e:
        logger.error(f"Error extracting HTML from PDF: {str(e)}")
        plain_text = extract_text_from_pdf(file)
        html_with_breaks = plain_text.replace('\n', "</div><div class='paragraph'>")
        return f'<div class="docx-content"><div class="paragraph">{html_with_breaks}</div></div>'

def extract_text_from_pdf(file):
    try:
        logger.info(f"Starting PDF processing for file: {file.name}, size: {file.size}")
        
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as pdf_temp:
            for chunk in file.chunks():
                pdf_temp.write(chunk)
            pdf_path = pdf_temp.name
            
        docx_path = pdf_path.replace('.pdf', '.docx')
        
        try:
            logger.info("Converting PDF to DOCX")
            cv = Converter(pdf_path)
            cv.convert(docx_path)
            cv.close()
            logger.info("PDF to DOCX conversion completed")
            
            logger.info("Extracting text from DOCX")
            doc = docx.Document(docx_path)
            text = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            full_text = '\n'.join(text)
            logger.info(f"Successfully extracted text, length: {len(full_text)}")
            
            if is_scanned_pdf(full_text):
                logger.error("PDF appears to be a scanned document")
                raise ValueError("Этот PDF-файл является сканированным документом. Пожалуйста, загрузите PDF с текстовым содержимым.")
            
            if not full_text.strip():
                logger.error("No text content extracted from PDF")
                raise ValueError("Не удалось извлечь текст из PDF файла. Возможно, это сканированный документ или файл поврежден.")
            
            return full_text
            
        finally:
            try:
                os.unlink(pdf_path)
                os.unlink(docx_path)
            except Exception as e:
                logger.warning(f"Failed to delete temporary files: {str(e)}")
        
    except ValueError as e:
        raise
    except Exception as e:
        logger.error(f"Unexpected error processing PDF: {str(e)}", exc_info=True)
        raise ValueError(f"Ошибка при обработке PDF файла: {str(e)}")
    
def extract_html_from_docx(file):
    """
    Извлекает HTML из файла DOCX с сохранением форматирования
    
    Args:
        file: Django UploadedFile объект
        
    Returns:
        str: HTML-представление документа
    """
    file.seek(0)
    doc = Document(file)
    html_parts = []
    
    html_parts.append('<div class="docx-content">')
    
    # Переменные для отслеживания списков
    current_list_type = None
    current_list_level = 0
    
    for para in doc.paragraphs:
        # Определяем тип списка
        list_info = get_list_info(para)
        
        if not para.text.strip():
            # Пустой параграф может завершать список
            if current_list_type:
                html_parts.append(f'</{current_list_type}>')
                current_list_type = None
                current_list_level = 0
            html_parts.append('<div class="paragraph empty-paragraph">&nbsp;</div>')
            continue
        
        # Обрабатываем списки
        if list_info['is_list']:
            new_list_type = 'ol' if list_info['is_numbered'] else 'ul'
            new_level = list_info['level']
            
            # Если начинается новый список или меняется тип
            if current_list_type != new_list_type or current_list_level != new_level:
                # Закрываем предыдущий список если есть
                if current_list_type:
                    html_parts.append(f'</{current_list_type}>')
                
                # Открываем новый список
                html_parts.append(f'<{new_list_type} class="docx-list level-{new_level}">')
                current_list_type = new_list_type
                current_list_level = new_level
            
            # Добавляем элемент списка
            html_parts.append('<li>')
            
            # Обрабатываем runs
            for run in para.runs:
                if not run.text:
                    continue
                    
                text = run.text.replace('<', '&lt;').replace('>', '&gt;').replace('&', '&amp;')
                
                # Применяем форматирование
                if run.bold:
                    text = f'<strong>{text}</strong>'
                if run.italic:
                    text = f'<em>{text}</em>'
                if run.underline:
                    text = f'<u>{text}</u>'
                    
                html_parts.append(text)
            
            html_parts.append('</li>')
            
        else:
            # Обычный параграф - закрываем список если был открыт
            if current_list_type:
                html_parts.append(f'</{current_list_type}>')
                current_list_type = None
                current_list_level = 0
            
            # Определяем выравнивание
            alignment = 'left'
            if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                alignment = 'center'
            elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                alignment = 'right'
            elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                alignment = 'justify'
            
            html_parts.append(f'<div class="paragraph" style="text-align: {alignment};">')
            
            # Обрабатываем runs (фрагменты с одинаковым форматированием)
            for run in para.runs:
                if not run.text:
                    continue
                    
                text = run.text.replace('<', '&lt;').replace('>', '&gt;').replace('&', '&amp;')
                
                # Применяем форматирование
                if run.bold:
                    text = f'<strong>{text}</strong>'
                if run.italic:
                    text = f'<em>{text}</em>'
                if run.underline:
                    text = f'<u>{text}</u>'
                    
                html_parts.append(text)
            
            html_parts.append('</div>')
    
    # Закрываем последний список если остался открытым
    if current_list_type:
        html_parts.append(f'</{current_list_type}>')
    
    # Обрабатываем таблицы
    for table in doc.tables:
        html_parts.append('<table class="docx-table">')
        for row in table.rows:
            html_parts.append('<tr>')
            for cell in row.cells:
                cell_html = []
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        cell_html.append(paragraph.text.replace('<', '&lt;').replace('>', '&gt;').replace('&', '&amp;'))
                html_parts.append(f'<td>{"<br>".join(cell_html)}</td>')
            html_parts.append('</tr>')
        html_parts.append('</table>')
    
    html_parts.append('</div>')
    
    return ''.join(html_parts)

def extract_text_from_docx(file):
    file.seek(0)
    doc = docx.Document(file)
    full_text = []
    
    for para in doc.paragraphs:
        full_text.append(para.text)
        
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text.append(paragraph.text)
    
    return '\n'.join(full_text)


